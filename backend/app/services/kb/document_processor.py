"""
Document Processor Service (MVP)
Extracts text from documents and chunks them for indexing
"""
import re
import uuid
from typing import List, Optional, Dict, Any
from pathlib import Path

# PDF processing
try:
    import pdfplumber
    import pypdf
except ImportError:
    pdfplumber = None
    pypdf = None

# DOCX processing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Tokenization
try:
    import tiktoken
except ImportError:
    tiktoken = None


class DocumentChunk:
    """Represents a text chunk from a document"""
    def __init__(
        self,
        text: str,
        chunk_index: int,
        metadata: Optional[Dict[str, Any]] = None
    ):
        self.id = str(uuid.uuid4())
        self.text = text
        self.chunk_index = chunk_index
        self.metadata = metadata or {}
        self.token_count = self._count_tokens(text)

    def _count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken"""
        if tiktoken:
            try:
                encoding = tiktoken.get_encoding("cl100k_base")
                return len(encoding.encode(text))
            except Exception:
                pass
        # Fallback: approximate 1 token = 4 chars
        return len(text) // 4

    def to_dict(self) -> Dict[str, Any]:
        return {
            'id': self.id,
            'text': self.text,
            'chunk_index': self.chunk_index,
            'token_count': self.token_count,
            'metadata': self.metadata
        }


class DocumentProcessor:
    """Process documents and extract text chunks"""

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        max_chunks: int = 1000
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_chunks = max_chunks

    async def process_document(
        self,
        file_path: str,
        document_id: str,
        project_id: str
    ) -> List[DocumentChunk]:
        """
        Process a document and return chunks

        Args:
            file_path: Path to document file
            document_id: UUID of document
            project_id: UUID of project

        Returns:
            List of DocumentChunk objects
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        # Extract text based on file type
        file_ext = path.suffix.lower()

        if file_ext == '.pdf':
            text = await self._extract_pdf(path)
        elif file_ext in ['.docx', '.doc']:
            text = await self._extract_docx(path)
        elif file_ext in ['.md', '.txt']:
            text = await self._extract_text(path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Clean text
        text = self._clean_text(text)

        # Split into chunks
        chunks = self._chunk_text(
            text=text,
            metadata={
                'document_id': document_id,
                'project_id': project_id,
                'file_type': file_ext,
                'file_name': path.name
            }
        )

        # Limit chunks
        if len(chunks) > self.max_chunks:
            chunks = chunks[:self.max_chunks]

        return chunks

    async def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF using pdfplumber (better accuracy)"""
        if not pdfplumber:
            raise ImportError("pdfplumber not installed")

        text_parts = []

        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            # Fallback to pypdf if pdfplumber fails
            if pypdf:
                return await self._extract_pdf_pypdf(path)
            raise e

        return '\n\n'.join(text_parts)

    async def _extract_pdf_pypdf(self, path: Path) -> str:
        """Fallback PDF extraction using pypdf"""
        if not pypdf:
            raise ImportError("pypdf not installed")

        text_parts = []

        with open(path, 'rb') as file:
            reader = pypdf.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return '\n\n'.join(text_parts)

    async def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX"""
        if not DocxDocument:
            raise ImportError("python-docx not installed")

        doc = DocxDocument(path)

        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return '\n\n'.join(text_parts)

    async def _extract_text(self, path: Path) -> str:
        """Extract plain text or markdown"""
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()

    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)

        # Trim
        text = text.strip()

        return text

    def _chunk_text(
        self,
        text: str,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks

        Uses token-based chunking with overlap for better context preservation
        """
        chunks = []

        if not tiktoken:
            # Fallback: character-based chunking
            return self._chunk_text_chars(text, metadata)

        # Token-based chunking
        try:
            encoding = tiktoken.get_encoding("cl100k_base")
            tokens = encoding.encode(text)

            chunk_index = 0
            start = 0

            while start < len(tokens):
                # Get chunk of tokens
                end = min(start + self.chunk_size, len(tokens))
                chunk_tokens = tokens[start:end]

                # Decode back to text
                chunk_text = encoding.decode(chunk_tokens)

                # Create chunk
                chunk = DocumentChunk(
                    text=chunk_text,
                    chunk_index=chunk_index,
                    metadata={
                        **metadata,
                        'start_token': start,
                        'end_token': end
                    }
                )
                chunks.append(chunk)

                # Move start position with overlap
                start += self.chunk_size - self.chunk_overlap
                chunk_index += 1

        except Exception:
            # Fallback to character-based
            return self._chunk_text_chars(text, metadata)

        return chunks

    def _chunk_text_chars(
        self,
        text: str,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Fallback: character-based chunking (approximate 4 chars = 1 token)"""
        char_chunk_size = self.chunk_size * 4
        char_overlap = self.chunk_overlap * 4

        chunks = []
        chunk_index = 0
        start = 0

        while start < len(text):
            end = min(start + char_chunk_size, len(text))
            chunk_text = text[start:end]

            chunk = DocumentChunk(
                text=chunk_text,
                chunk_index=chunk_index,
                metadata={
                    **metadata,
                    'start_char': start,
                    'end_char': end
                }
            )
            chunks.append(chunk)

            start += char_chunk_size - char_overlap
            chunk_index += 1

        return chunks


# Singleton instance
_processor_instance = None

def get_document_processor() -> DocumentProcessor:
    """Get singleton DocumentProcessor instance"""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = DocumentProcessor(
            chunk_size=512,
            chunk_overlap=50,
            max_chunks=1000
        )
    return _processor_instance
